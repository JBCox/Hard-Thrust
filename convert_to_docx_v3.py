"""
Convert FULL_MANUSCRIPT.txt to properly formatted .docx for KDP
VERSION 4: Supports both Print (with page numbers) and Kindle (no page numbers)

Usage:
    python convert_to_docx_v3.py --format print   # Creates HARD_THRUST_KDP_PRINT.docx
    python convert_to_docx_v3.py --format kindle  # Creates HARD_THRUST_KINDLE.docx
    python convert_to_docx_v3.py                  # Default: print

Print version (HARD_THRUST_KDP_PRINT.docx):
- Page size: 6" x 9"
- Margins: 0.75" inside, 0.5" outside/top/bottom (mirror margins for binding)
- Two sections: Front matter (no footers) + Body (with page numbers)
- Page numbers in footer: odd pages=right aligned, even pages=left aligned
- Numbering starts at 1 on Chapter 1 (not front matter)
- Section break after title page separates front matter from body

Kindle version (HARD_THRUST_KINDLE.docx):
- Page size: 6" x 9" (ignored by Kindle, but needed for Word display)
- Margins: 0.5" all sides (no mirror margins - ebooks don't have binding)
- No page numbers (Kindle handles pagination dynamically)
- Simple page breaks (no section breaks needed)

Common formatting (both versions):
- Font: Times New Roman 11pt
- Line spacing: 1.25
- First-line indent: 0.3" (except first paragraph after breaks)
- Text alignment: Justified
- Chapter headings: 16pt bold, centered
- Scene breaks: * * * centered
"""
import re
import argparse
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def add_page_number_field(paragraph):
    """Insert a PAGE field into a paragraph (displays current page number)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Style the run
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    return run


def setup_body_section_footers(section):
    """Configure odd/even page footers with page numbers for the body section."""
    # Enable different odd/even headers/footers
    sectPr = section._sectPr

    # Add titlePg element to suppress header/footer on first page of section (optional)
    # We'll skip this since we want page numbers on all body pages

    # Create footer for odd pages (right-hand pages) - number on RIGHT
    odd_footer = section.footer
    odd_footer.is_linked_to_previous = False
    odd_para = odd_footer.paragraphs[0] if odd_footer.paragraphs else odd_footer.add_paragraph()
    odd_para.clear()
    odd_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number_field(odd_para)

    # Create footer for even pages (left-hand pages) - number on LEFT
    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    even_para = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
    even_para.clear()
    even_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_page_number_field(even_para)


def restart_page_numbering(section, start=1):
    """Restart page numbering at a specific number for this section."""
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def enable_odd_even_footers(doc):
    """Enable different odd/even page headers/footers in document settings."""
    settings = doc.settings
    # Access the document settings element
    settingsElement = settings.element
    evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
    settingsElement.append(evenAndOddHeaders)

def convert_manuscript(output_format='print'):
    """
    Convert manuscript to DOCX.

    Args:
        output_format: 'print' for KDP Print (with page numbers) or
                      'kindle' for Kindle ebook (no page numbers)
    """
    # Read the source file
    with open('FULL_MANUSCRIPT.txt', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # For print version, enable odd/even footers at document level
    if output_format == 'print':
        enable_odd_even_footers(doc)

    # ============================================
    # SET UP PAGE SIZE AND MARGINS FOR 6x9
    # ============================================
    section = doc.sections[0]

    # Page size: 6" x 9"
    section.page_width = Inches(6)
    section.page_height = Inches(9)

    if output_format == 'kindle':
        # Kindle: Equal margins on all sides, no mirror margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.gutter = Inches(0)
    else:
        # Print: Margins for 6x9 with proper gutter for 328+ pages
        # KDP requires: 0.625" gutter minimum for 328 pages
        # With mirror margins: left = inside (gutter), right = outside
        section.left_margin = Inches(0.75)   # Inside/gutter margin
        section.right_margin = Inches(0.5)   # Outside margin
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.gutter = Inches(0)

        # Enable mirror margins for proper book layout
        # This swaps left/right on even pages so gutter is always on spine side
        sectPr = section._sectPr
        mirrorMargins = OxmlElement('w:mirrorMargins')
        sectPr.append(mirrorMargins)

    # ============================================
    # SET UP STYLES
    # ============================================
    styles = doc.styles

    # Modify Normal style - for regular body paragraphs
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)  # Industry standard for 6x9
    normal_style.paragraph_format.line_spacing = 1.25  # Industry standard
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.first_line_indent = Inches(0.3)  # Modern standard
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified text

    # Create "No Indent" style for first paragraphs after breaks
    try:
        no_indent_style = styles.add_style('NoIndent', 1)  # 1 = paragraph style
    except:
        no_indent_style = styles['NoIndent']
    no_indent_style.base_style = normal_style
    no_indent_style.font.name = 'Times New Roman'
    no_indent_style.font.size = Pt(11)
    no_indent_style.paragraph_format.first_line_indent = Inches(0)  # No indent for first para
    no_indent_style.paragraph_format.line_spacing = 1.25
    no_indent_style.paragraph_format.space_after = Pt(0)
    no_indent_style.paragraph_format.space_before = Pt(0)
    no_indent_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Modify Heading 1 for chapters
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(36)  # More space before chapter titles
    h1_style.paragraph_format.space_after = Pt(18)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_style.paragraph_format.first_line_indent = Inches(0)

    # Modify Heading 2 for date/POV headers
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = False
    h2_style.font.italic = True
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(12)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2_style.paragraph_format.first_line_indent = Inches(0)

    # Modify Heading 3 for sub-section headers
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(11)
    h3_style.font.bold = False
    h3_style.font.italic = True
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h3_style.paragraph_format.first_line_indent = Inches(0)

    # ============================================
    # PROCESS MANUSCRIPT CONTENT
    # ============================================
    lines = content.split('\n')

    i = 0
    in_front_matter = True
    next_para_no_indent = False
    body_section_created = False  # Track whether we've created the body section

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines at start
        if not line:
            i += 1
            continue

        # Skip decorative lines
        if line.startswith('====='):
            i += 1
            continue

        # ---- COPYRIGHT PAGE ----
        if line.startswith('Copyright'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(6)
            i += 1

            # Get rest of copyright text
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith('=====') or line == 'CONTENT NOTE':
                    break
                if line:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.space_after = Pt(6)
                i += 1

            doc.add_page_break()
            continue

        # ---- CONTENT NOTE ----
        if line == 'CONTENT NOTE':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('CONTENT NOTE')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0)
            i += 1

            # Get content note text
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith('=====') or line.startswith('#') or line == 'HARD THRUST':
                    break
                if line:
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Inches(0)
                i += 1

            doc.add_page_break()
            continue

        # ---- TITLE PAGE ----
        if line == 'HARD THRUST' and in_front_matter:
            # Add some vertical space
            for _ in range(3):
                doc.add_paragraph()

            # Title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('HARD THRUST')
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(6)

            i += 1
            # Look for subtitle or author
            while i < len(lines) and not lines[i].strip():
                i += 1

            # Subtitle if present
            if i < len(lines) and 'A Novel' in lines[i]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('A Novel')
                run.font.size = Pt(14)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                p.paragraph_format.first_line_indent = Inches(0)
                i += 1
                while i < len(lines) and not lines[i].strip():
                    i += 1

            # Author
            if i < len(lines) and 'Joshua Cox' in lines[i]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                run = p.add_run('JOSHUA COX')
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
                p.paragraph_format.first_line_indent = Inches(0)
                i += 1
                while i < len(lines) and not lines[i].strip():
                    i += 1

            # Publisher if present
            if i < len(lines) and 'Notus' in lines[i]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(36)
                run = p.add_run(lines[i].strip())
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                p.paragraph_format.first_line_indent = Inches(0)
                i += 1

            if output_format == 'print':
                # For print: Add section break to create body section
                # This allows page numbers to start at 1 in the body
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)

                # Copy page size and margins to new section
                new_section.page_width = Inches(6)
                new_section.page_height = Inches(9)
                new_section.left_margin = Inches(0.75)
                new_section.right_margin = Inches(0.5)
                new_section.top_margin = Inches(0.5)
                new_section.bottom_margin = Inches(0.5)
                new_section.gutter = Inches(0)

                # Enable mirror margins for body section
                sectPr = new_section._sectPr
                mirrorMargins = OxmlElement('w:mirrorMargins')
                sectPr.append(mirrorMargins)

                # Set up page number footers for body section
                setup_body_section_footers(new_section)
                restart_page_numbering(new_section, start=1)
                body_section_created = True
            else:
                # For Kindle: just use page break
                doc.add_page_break()
            continue

        # ---- CHAPTER HEADING ----
        if line.startswith('# '):
            in_front_matter = False
            chapter_title = line[2:].strip()

            # Add page break before new chapters (except first)
            # Skip if we just created a section break (body_section_created)
            # because the section break already started a new page
            if len(doc.paragraphs) > 10 and not body_section_created:
                doc.add_page_break()

            # Reset flag after first chapter
            body_section_created = False

            # Add some vertical space at top of chapter
            for _ in range(2):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)

            doc.add_heading(chapter_title, level=1)
            next_para_no_indent = True
            i += 1
            continue

        # ---- SCENE/POV HEADING ----
        if line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            next_para_no_indent = True
            i += 1
            continue

        # ---- SUB-SECTION HEADING ----
        if line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)
            next_para_no_indent = True
            i += 1
            continue

        # ---- SCENE BREAK ----
        if line == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('* * *')
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0)
            next_para_no_indent = True
            i += 1
            continue

        # ---- ABOUT THE AUTHOR (Back Matter) ----
        if line == 'ABOUT THE AUTHOR':
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('ABOUT THE AUTHOR')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(18)
            i += 1

            # Get about text
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith('=====') or line == 'A NOTE FROM THE AUTHOR':
                    break
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(6)
                i += 1
            continue

        # ---- NOTE FROM AUTHOR (Back Matter) ----
        if line == 'A NOTE FROM THE AUTHOR':
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('A NOTE FROM THE AUTHOR')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(18)
            i += 1

            # Get note text
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith('====='):
                    break
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(6)
                i += 1
            continue

        # ---- REGULAR PARAGRAPH ----
        if line and not in_front_matter:
            p = doc.add_paragraph()

            # Apply appropriate style
            if next_para_no_indent:
                p.style = 'NoIndent'
                next_para_no_indent = False
            else:
                p.style = 'Normal'

            # Parse bold (**text**) and italics (*text*)
            def add_formatted_text(paragraph, text):
                # Split by bold first (**text**)
                bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for bold_part in bold_parts:
                    if bold_part.startswith('**') and bold_part.endswith('**'):
                        run = paragraph.add_run(bold_part[2:-2])
                        run.bold = True
                    else:
                        # Check for italics within this part
                        italic_parts = re.split(r'(\*[^*]+\*)', bold_part)
                        for italic_part in italic_parts:
                            if italic_part.startswith('*') and italic_part.endswith('*'):
                                run = paragraph.add_run(italic_part[1:-1])
                                run.italic = True
                            else:
                                if italic_part:
                                    paragraph.add_run(italic_part)

            add_formatted_text(p, line)

        i += 1

    # ============================================
    # SAVE DOCUMENT
    # ============================================
    if output_format == 'print':
        output_file = 'HARD_THRUST_KDP_PRINT.docx'
    else:
        output_file = 'HARD_THRUST_KINDLE.docx'

    doc.save(output_file)

    print(f'Created: {output_file}')
    print(f'Format: {output_format.upper()}')
    print(f'Document has {len(doc.paragraphs)} paragraphs')
    print(f'\nFormatting Applied:')
    print(f'  Page size: 6" x 9"')
    if output_format == 'print':
        print(f'  Margins: 0.75" inside, 0.5" outside/top/bottom (mirror margins)')
        print(f'  Page numbers: Yes (starting at 1 on Chapter 1)')
        print(f'  Footer: Odd pages=right, Even pages=left')
    else:
        print(f'  Margins: 0.5" all sides (no mirror margins)')
        print(f'  Page numbers: No')
    print(f'  Font: Times New Roman 11pt')
    print(f'  Line spacing: 1.25')
    print(f'  First line indent: 0.3"')
    print(f'  Text alignment: Justified')

    # Count style usage
    normal_count = 0
    no_indent_count = 0
    for para in doc.paragraphs:
        if para.style.name == 'Normal':
            normal_count += 1
        elif para.style.name == 'NoIndent':
            no_indent_count += 1

    print(f'\nParagraph counts:')
    print(f'  Normal style: {normal_count}')
    print(f'  NoIndent style: {no_indent_count}')

    return output_file


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Convert FULL_MANUSCRIPT.txt to formatted DOCX for KDP'
    )
    parser.add_argument(
        '--format',
        choices=['print', 'kindle'],
        default='print',
        help='Output format: print (with page numbers) or kindle (no page numbers)'
    )
    args = parser.parse_args()

    convert_manuscript(args.format)
